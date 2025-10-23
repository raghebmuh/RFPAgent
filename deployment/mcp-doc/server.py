"""
MCP-Doc Server: FastMCP server for Word document generation
Provides tools for creating, editing, and managing .docx files with Arabic RFP template support
"""

import os
import sys
import uuid
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from fastmcp import FastMCP

# Add application templates to path
sys.path.insert(0, '/app/application/templates')
try:
    from rfp_template_ar import RFPTemplateKSA, Section as RFPSection, TableSpec
except ImportError:
    RFPTemplateKSA = None
    RFPSection = None
    TableSpec = None
    print("Warning: Arabic RFP template not available")

# Initialize logging
logger = logging.getLogger(__name__)

# Initialize FastMCP server
mcp = FastMCP("MCP-Doc Server")

# Document storage
DOCUMENTS_DIR = Path("/app/documents")
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

# In-memory document store (doc_id -> Document object)
active_documents: Dict[str, Document] = {}
document_metadata: Dict[str, dict] = {}


def replace_placeholder_in_paragraph(paragraph, placeholder_pattern: str, replacement_value: str):
    """
    Replace placeholder in paragraph while preserving ALL formatting, RTL, and Arabic font.

    Args:
        paragraph: The docx paragraph object
        placeholder_pattern: The placeholder to find (e.g., "{{tender_name}}")
        replacement_value: The value to replace it with

    Returns:
        bool: True if replacement was made, False otherwise
    """
    # Check if placeholder exists in full paragraph text
    full_text = paragraph.text
    if placeholder_pattern not in full_text:
        return False

    # Replace the placeholder in the full text
    new_text = full_text.replace(placeholder_pattern, replacement_value)

    # Store the formatting properties of the first run
    if paragraph.runs:
        first_run = paragraph.runs[0]

        # Capture all formatting properties
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
        font_underline = first_run.font.underline
        font_color = first_run.font.color.rgb if first_run.font.color.rgb else None

        # Capture complex script font (for Arabic)
        cs_font = None
        if first_run._element.rPr is not None:
            rFonts = first_run._element.rPr.rFonts
            if rFonts is not None:
                cs_font = rFonts.get(qn('w:cs'))

        # Clear all runs
        for run in paragraph.runs:
            run.text = ""

        # Set the new text in the first run
        first_run.text = new_text

        # Restore all formatting
        if font_name:
            first_run.font.name = font_name
        if font_size:
            first_run.font.size = font_size
        if font_bold is not None:
            first_run.font.bold = font_bold
        if font_italic is not None:
            first_run.font.italic = font_italic
        if font_underline is not None:
            first_run.font.underline = font_underline
        if font_color:
            first_run.font.color.rgb = font_color

        # Restore complex script font for Arabic
        if cs_font or font_name:
            if first_run._element.rPr is None:
                first_run._element.get_or_add_rPr()
            rFonts = first_run._element.rPr.get_or_add_rFonts()
            rFonts.set(qn('w:cs'), cs_font or font_name or 'Sakkal Majalla')
            rFonts.set(qn('w:ascii'), font_name or 'Sakkal Majalla')
            rFonts.set(qn('w:hAnsi'), font_name or 'Sakkal Majalla')
    else:
        # No runs existed, add the text directly
        paragraph.text = new_text

    return True


def replace_placeholder_with_formatted_content(paragraph, placeholder_pattern: str, replacement_value: str, doc: Document):
    """
    Replace placeholder with multi-paragraph content while preserving formatting.
    For multi-line content, this creates additional paragraphs after the current one.

    Args:
        paragraph: The docx paragraph object
        placeholder_pattern: The placeholder to find
        replacement_value: The value to replace (can contain newlines)
        doc: The document object to add new paragraphs

    Returns:
        bool: True if replacement was made
    """
    full_text = paragraph.text
    if placeholder_pattern not in full_text:
        return False

    # Check if this is multi-line content
    lines = replacement_value.split('\n')

    if len(lines) == 1:
        # Single line - use simple replacement
        return replace_placeholder_in_paragraph(paragraph, placeholder_pattern, replacement_value)

    # Multi-line content - replace first line and add subsequent lines as new paragraphs
    # Replace with first line
    first_line = lines[0]
    replace_placeholder_in_paragraph(paragraph, placeholder_pattern, first_line)

    # Capture formatting from the original paragraph
    if paragraph.runs:
        original_run = paragraph.runs[0]
        font_name = original_run.font.name
        font_size = original_run.font.size
        font_bold = original_run.font.bold
        cs_font = None
        if original_run._element.rPr is not None:
            rFonts = original_run._element.rPr.rFonts
            if rFonts is not None:
                cs_font = rFonts.get(qn('w:cs'))
    else:
        font_name = 'Sakkal Majalla'
        font_size = Pt(16)
        font_bold = False
        cs_font = 'Sakkal Majalla'

    # Get paragraph's parent element to insert after current paragraph
    parent = paragraph._element.getparent()
    paragraph_index = parent.index(paragraph._element)

    # Add remaining lines as new paragraphs right after current one
    for i, line in enumerate(lines[1:], 1):
        if not line.strip():
            continue

        # Create new paragraph element
        new_p = OxmlElement('w:p')
        parent.insert(paragraph_index + i, new_p)

        # Create paragraph object from element
        new_para = Paragraph(new_p, paragraph._parent)

        # Add run with text
        run = new_para.add_run(line)

        # Apply formatting
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_bold is not None:
            run.font.bold = font_bold

        # Apply Arabic font
        if run._element.rPr is None:
            run._element.get_or_add_rPr()
        rFonts = run._element.rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), cs_font or font_name or 'Sakkal Majalla')
        rFonts.set(qn('w:ascii'), font_name or 'Sakkal Majalla')

        # Set RTL
        set_rtl_paragraph(new_para)

    return True


def set_rtl_paragraph(paragraph):
    """Set paragraph direction to RTL (Right-to-Left) for Arabic text."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    return paragraph


def set_arabic_font(run, font_name="Sakkal Majalla", font_size=16):
    """Set Arabic font for a text run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    # Set font for complex scripts (Arabic, Hebrew, etc.)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)
    return run


@mcp.tool()
def create_rfp_document(
    title: str,
    project_name: str,
    company_name: Optional[str] = None,
    date: Optional[str] = None
) -> dict:
    """
    Create a new RFP (Request for Proposal) Word document with standard structure.

    Args:
        title: The title of the RFP document
        project_name: Name of the project for which RFP is being created
        company_name: Optional company name issuing the RFP
        date: Optional date (defaults to current date)

    Returns:
        dict with doc_id, title, and initial structure
    """
    doc_id = str(uuid.uuid4())
    doc = Document()

    # Set up document properties
    core_properties = doc.core_properties
    core_properties.title = title
    core_properties.subject = f"RFP for {project_name}"
    core_properties.author = company_name or "RFPAgent"

    # Add title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add document info
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run(f"Project: {project_name}").bold = True
    if company_name:
        doc.add_paragraph(f"Issued by: {company_name}")
    doc.add_paragraph(f"Date: {date or datetime.now().strftime('%B %d, %Y')}")

    # Add page break
    doc.add_page_break()

    # Store document
    active_documents[doc_id] = doc
    document_metadata[doc_id] = {
        "doc_id": doc_id,
        "title": title,
        "project_name": project_name,
        "created_at": datetime.now().isoformat(),
        "sections": []
    }

    # AUTO-SAVE: Save document to disk immediately after creation
    safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_', 'ا', 'ب', 'ت', 'ث', 'ج', 'ح', 'خ', 'د', 'ذ', 'ر', 'ز', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ك', 'ل', 'م', 'ن', 'ه', 'و', 'ي') else '_' for c in title)
    safe_title = safe_title.replace(' ', '_')[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"RFP_{safe_title}_{doc_id[:8]}_{timestamp}.docx"
    file_path = DOCUMENTS_DIR / file_name
    doc.save(str(file_path))

    # Update metadata with file information
    document_metadata[doc_id]["file_path"] = str(file_path)
    document_metadata[doc_id]["file_name"] = file_name
    document_metadata[doc_id]["saved_at"] = datetime.now().isoformat()

    # Construct download URL
    api_host = os.getenv("API_HOST", "http://localhost:7091")
    download_url = f"{api_host}/api/documents/download/{doc_id}"

    return {
        "success": True,
        "doc_id": doc_id,
        "title": title,
        "file_path": str(file_path),
        "file_name": file_name,
        "download_url": download_url,
        "message": f"RFP document '{title}' created and saved successfully. Download at: {download_url}"
    }


@mcp.tool()
def create_rfp_from_template(
    template_name: str,
    title: str,
    project_name: str,
    technical_organization_name: Optional[str] = None,
    tender_number: Optional[str] = None,
    project_description: Optional[str] = None,
    scope_of_work: Optional[str] = None,
    budget: Optional[str] = None,
    duration: Optional[str] = None,
    date: Optional[str] = None
) -> dict:
    """
    Create an RFP document by loading and filling a reference template document.

    This tool loads an existing reference RFP .docx file, fills in the user's
    project details, and saves it as a new document.

    Args:
        template_name: Name of the reference RFP file to use as template (e.g., "كراسة تطبيق أتمتة خدمات تقنية المعلومات للمركز (المرحلة الأولى).docx")
        title: The title for the new RFP document
        project_name: Name of the project
        technical_organization_name: Optional name of the government entity
        tender_number: Optional tender/RFP number
        project_description: Optional detailed project description
        scope_of_work: Optional scope of work details
        budget: Optional budget information
        duration: Optional project duration
        date: Optional date (defaults to current date)

    Returns:
        dict with doc_id, title, file_path, and success status
    """
    doc_id = str(uuid.uuid4())
    reference_dir = Path("/app/application/templates/RFPs")
    template_path = reference_dir / template_name

    if not template_path.exists():
        return {
            "success": False,
            "error": f"Template file '{template_name}' not found in templates directory"
        }

    if template_path.suffix.lower() not in ['.docx']:
        return {
            "success": False,
            "error": "Only .docx templates are supported"
        }

    try:
        from docx import Document as DocxDocument

        # Load the template document
        doc = DocxDocument(str(template_path))

        # Create a mapping of placeholders to replacement values
        replacements = {
            "[اسم المشروع]": project_name,
            "[اسم المنافسة]": project_name,
            "[اسم الجهة]": technical_organization_name or "[اسم الجهة]",
            "[الجهة الحكومية]": technical_organization_name or "[اسم الجهة]",
            "[رقم الكراسة]": tender_number or "[رقم الكراسة]",
            "[رقم المنافسة]": tender_number or "[رقم الكراسة]",
            "[التاريخ]": date or datetime.now().strftime('%Y/%m/%d'),
            "[وصف المشروع]": project_description or "[وصف المشروع]",
            "[نطاق العمل]": scope_of_work or "[نطاق العمل]",
            "[الميزانية]": budget or "[الميزانية]",
            "[المدة]": duration or "[المدة]",
        }

        # Replace text in all paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    # Replace in inline text
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)

        # Replace text in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in replacements.items():
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, replacement)

        # Update document properties
        core_properties = doc.core_properties
        core_properties.title = title
        core_properties.subject = f"كراسة الشروط والمواصفات - {project_name}"
        core_properties.author = technical_organization_name or "RFPAgent"

        # Store document in memory
        active_documents[doc_id] = doc

        # Extract sections for metadata (simplified - just count paragraphs with heading styles)
        sections = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                sections.append({
                    "code": f"S{i}",
                    "title": paragraph.text,
                    "heading": paragraph.text,
                    "level": level
                })

        # Save metadata
        document_metadata[doc_id] = {
            "doc_id": doc_id,
            "title": title,
            "project_name": project_name,
            "technical_organization_name": technical_organization_name,
            "tender_number": tender_number,
            "created_at": datetime.now().isoformat(),
            "language": "ar",
            "rtl": True,
            "template_used": template_name,
            "sections": sections[:50]  # Limit to first 50 sections
        }

        # AUTO-SAVE: Save document to disk immediately with Arabic naming format
        safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_', 'ا', 'ب', 'ت', 'ث', 'ج', 'ح', 'خ', 'د', 'ذ', 'ر', 'ز', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ك', 'ل', 'م', 'ن', 'ه', 'و', 'ي') else '_' for c in title)
        safe_title = safe_title.replace(' ', '_')[:50]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"RFP_{safe_title}_{doc_id[:8]}_{timestamp}.docx"
        file_path = DOCUMENTS_DIR / file_name
        doc.save(str(file_path))

        # Update metadata with file information
        document_metadata[doc_id]["file_path"] = str(file_path)
        document_metadata[doc_id]["file_name"] = file_name
        document_metadata[doc_id]["saved_at"] = datetime.now().isoformat()

        # Construct download URL
        api_host = os.getenv("API_HOST", "http://localhost:7091")
        download_url = f"{api_host}/api/documents/download/{doc_id}"

        return {
            "success": True,
            "doc_id": doc_id,
            "title": title,
            "file_path": str(file_path),
            "file_name": file_name,
            "download_url": download_url,
            "template_used": template_name,
            "sections_count": len(sections),
            "message": f"تم إنشاء كراسة الشروط '{title}' بنجاح من القالب '{template_name}' وحفظها تلقائياً. التحميل: {download_url}"
        }

    except Exception as e:
        logger.error(f"Error creating RFP from template: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Failed to create RFP from template: {str(e)}"
        }

 
@mcp.tool()
def create_arabic_rfp_document(
    placeholders: dict,
    title: str,
    tender_name: str,
    technical_organization_name: Optional[str] = None,
    tender_number: Optional[str] = None,
    date: Optional[str] = None,
    conversation_id: Optional[str] = None
) -> dict:
    """
    Fill the RFP template with placeholder data and generate a downloadable document with full KSA structure.

    This tool fills the specific template file at inputs/templates/rfp_template_with_placeholders.docx
    with the provided placeholder values and generates a ready-to-download RFP document.

    Args:
        placeholders: Dictionary of placeholder names and their values (e.g., {"tender_name": "Project XYZ", "tender_number": "TC-2025-001"})
        title: The title of the RFP document (Arabic)
        tender_name: Name of the tender/project for the document title (Arabic)
        technical_organization_name: Optional government entity name (Arabic)
        tender_number: Optional tender/RFP number (Arabic)
        date: Optional date (defaults to current date in Arabic)
        conversation_id: Optional conversation ID to link the document (Arabic)

    Returns:
        dict with success status, doc_id, download_button, and file information (Arabic)
    """
    try:
        from docx import Document as DocxDocument

        # Template path (mounted from host)
        template_path = Path("/app/inputs/templates/rfp_template_with_placeholders.docx")

        if not template_path.exists():
            return {
                "success": False,
                "error": f"Template file not found at {template_path}. Please ensure the template is mounted correctly."
            }

        # Generate document ID
        doc_id = str(uuid.uuid4())

        # Load the template
        doc = DocxDocument(str(template_path))

        # Set up document properties
        core_properties = doc.core_properties
        core_properties.title = title
        core_properties.subject = f"كراسة الشروط والمواصفات - {tender_name}"
        core_properties.author = technical_organization_name or "RFPAgent"

        # Define ALL expected placeholders with defaults
        all_expected_placeholders = {
            "tender_name": tender_name,
            "tender_number": tender_number or "يرجى تحديد رقم المنافسة",
            "technical_organization_name": technical_organization_name or "الجهة الحكومية",
            "tender_purpose": "سيتم تحديده",
            "tender_documents_fees": "سيتم تحديده",
            "definition_department": "الإدارة المختصة",
            "project_scope": "سيتم تحديد نطاق العمل",
            "work_execution_method": "سيتم تحديد طريقة تنفيذ الأعمال",
            "work_program_phases": "سيتم تحديد مراحل البرنامج الزمني",
            "work_program_payment_method": "سيتم تحديد طريقة الدفع",
            "technical_inquiries_entity_name": technical_organization_name or "الجهة المختصة",
            "technical_inquiries_email": "inquiries@example.gov.sa",
            "technical_inquiries_alt_email": "سيتم تحديده",
            "technical_inquiries_duration": "5 أيام عمل",
            "bids_review_proposals": "سيتم تحديد معايير المراجعة",
            "purchase_reference": "سيتم تحديده",
            "supplier_samples_delivery_address": "سيتم تحديده",
            "samples_delivery_building": "سيتم تحديده",
            "samples_delivery_floor": "سيتم تحديده",
            "samples_delivery_room_or_department": "سيتم تحديده",
            "samples_delivery_time": "خلال ساعات العمل الرسمية"
        }

        # Merge user-provided placeholders with defaults
        final_placeholders = {**all_expected_placeholders, **placeholders}

        # Log placeholder summary
        logger.info(f"Starting placeholder replacement - Total: {len(final_placeholders)} placeholders")
        logger.info(f"User-provided: {len(placeholders)}, Using defaults: {len(all_expected_placeholders) - len(placeholders)}")

        # Replace placeholders in paragraphs (using {{placeholder}} format)
        replacements_made = 0
        for paragraph in doc.paragraphs:
            for placeholder_name, value in final_placeholders.items():
                placeholder_pattern = "{{" + placeholder_name + "}}"
                if replace_placeholder_in_paragraph(paragraph, placeholder_pattern, str(value)):
                    replacements_made += 1
                    logger.info(f"✓ Replaced {placeholder_pattern} with: {str(value)[:50]}...")

        logger.info(f"Replaced {replacements_made} placeholders in paragraphs")

        # Replace placeholders in tables (using {{placeholder}} format)
        table_replacements = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder_name, value in final_placeholders.items():
                            placeholder_pattern = "{{" + placeholder_name + "}}"
                            if replace_placeholder_in_paragraph(paragraph, placeholder_pattern, str(value)):
                                table_replacements += 1
                                logger.info(f"✓ Replaced {placeholder_pattern} in table cell")

        logger.info(f"Replaced {table_replacements} placeholders in tables")

        # Replace placeholders in headers and footers (using {{placeholder}} format)
        header_footer_replacements = 0
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                for placeholder_name, value in final_placeholders.items():
                    placeholder_pattern = "{{" + placeholder_name + "}}"
                    if replace_placeholder_in_paragraph(paragraph, placeholder_pattern, str(value)):
                        header_footer_replacements += 1
                        logger.info(f"✓ Replaced {placeholder_pattern} in header")

            # Footer
            for paragraph in section.footer.paragraphs:
                for placeholder_name, value in final_placeholders.items():
                    placeholder_pattern = "{{" + placeholder_name + "}}"
                    if replace_placeholder_in_paragraph(paragraph, placeholder_pattern, str(value)):
                        header_footer_replacements += 1
                        logger.info(f"✓ Replaced {placeholder_pattern} in footer")

        logger.info(f"Replaced {header_footer_replacements} placeholders in headers/footers")

        total_replacements = replacements_made + table_replacements + header_footer_replacements
        logger.info(f"✅ TOTAL REPLACEMENTS: {total_replacements} placeholders replaced successfully")
        logger.info(f"📝 All {len(final_placeholders)} expected placeholders have been processed")

        # Generate safe filename with doc_id for easy lookup (Arabic support)
        safe_tender_name = "".join(c if c.isalnum() or c in (' ', '-', '_', 'ا', 'ب', 'ت', 'ث', 'ج', 'ح', 'خ', 'د', 'ذ', 'ر', 'ز', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ك', 'ل', 'م', 'ن', 'ه', 'و', 'ي') else '_' for c in tender_name)
        safe_tender_name = safe_tender_name.replace(' ', '_')[:50]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Include first 8 chars of doc_id in filename for download lookup
        file_name = f"RFP_{safe_tender_name}_{doc_id[:8]}_{timestamp}.docx"

        # Save to documents directory
        file_path = DOCUMENTS_DIR / file_name
        doc.save(str(file_path))

        # Store document metadata
        active_documents[doc_id] = doc
        document_metadata[doc_id] = {
            "doc_id": doc_id,
            "title": f"RFP - {tender_name}",
            "tender_name": tender_name,
            "conversation_id": conversation_id,
            "created_at": datetime.now().isoformat(),
            "file_path": str(file_path),
            "file_name": file_name,
            "saved_at": datetime.now().isoformat(),
            "template_used": "rfp_template_with_placeholders.docx",
            "placeholders_filled": len(placeholders),
            "language": "ar",
            "rtl": True
        }

        # Construct download URL
        api_host = os.getenv("API_HOST", "http://localhost:7091")
        download_url = f"{api_host}/api/documents/download/{doc_id}"

        logger.info(f"Successfully filled RFP template. Doc ID: {doc_id}, File: {file_name}")

        return {
            "success": True,
            "doc_id": doc_id,
            "title": f"RFP - {tender_name}",
            "file_name": file_name,
            "file_path": str(file_path),
            "download_url": download_url,
            "placeholders_filled": len(placeholders),
            "template_used": "rfp_template_with_placeholders.docx",
            "message": f"تم إنشاء وثيقة RFP بنجاح باستخدام القالب المحدد. تم ملء {len(placeholders)} حقل. يمكنك تحميل الوثيقة من الرابط."
        }

    except Exception as e:
        logger.error(f"Error filling RFP template: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Failed to fill template: {str(e)}"
        }
        


@mcp.tool()
def add_section(
    doc_id: str,
    heading: str,
    content: str,
    level: int = 1
) -> dict:
    """
    Add a section with heading and content to the document.

    Args:
        doc_id: Document ID from create_rfp_document
        heading: Section heading text
        content: Section content (supports paragraphs separated by newlines)
        level: Heading level (1-4, default 1)

    Returns:
        dict with success status
    """
    if doc_id not in active_documents:
        return {"success": False, "error": f"Document {doc_id} not found"}

    doc = active_documents[doc_id]
    metadata = document_metadata.get(doc_id, {})
    is_rtl = metadata.get("rtl", False)

    # Add heading
    heading_para = doc.add_heading(heading, level=level)
    if is_rtl:
        set_rtl_paragraph(heading_para)
        for run in heading_para.runs:
            set_arabic_font(run, font_size=18 if level == 1 else 16)

    # Add content paragraphs
    paragraphs = content.strip().split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Check if it's a bullet point
            if para_text.strip().startswith('-') or para_text.strip().startswith('•'):
                # Remove bullet marker and add as list item
                text = para_text.strip().lstrip('-•').strip()
                para = doc.add_paragraph(text, style='List Bullet')
            elif para_text.strip()[0].isdigit() and '.' in para_text.strip()[:3]:
                # Numbered list
                text = para_text.strip().split('.', 1)[1].strip()
                para = doc.add_paragraph(text, style='List Number')
            else:
                para = doc.add_paragraph(para_text.strip())

            # Apply RTL formatting if needed
            if is_rtl:
                set_rtl_paragraph(para)
                for run in para.runs:
                    set_arabic_font(run)

    # Update metadata
    if doc_id in document_metadata:
        document_metadata[doc_id]["sections"].append({
            "heading": heading,
            "level": level,
            "content_length": len(content)
        })

    return {
        "success": True,
        "doc_id": doc_id,
        "section_added": heading,
        "message": f"Section '{heading}' added successfully"
    }


@mcp.tool()
def add_table(
    doc_id: str,
    rows: int,
    cols: int,
    data: Optional[List[List[str]]] = None,
    header_row: bool = True
) -> dict:
    """
    Add a table to the document.

    Args:
        doc_id: Document ID
        rows: Number of rows
        cols: Number of columns
        data: Optional 2D list of cell data
        header_row: Whether first row is a header

    Returns:
        dict with success status
    """
    if doc_id not in active_documents:
        return {"success": False, "error": f"Document {doc_id} not found"}

    doc = active_documents[doc_id]
    metadata = document_metadata.get(doc_id, {})
    is_rtl = metadata.get("rtl", False)

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    # Fill in data if provided
    if data:
        for i, row_data in enumerate(data[:rows]):
            for j, cell_data in enumerate(row_data[:cols]):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_data)

                # Apply RTL and Arabic font to table cells
                if is_rtl:
                    for paragraph in cell.paragraphs:
                        set_rtl_paragraph(paragraph)
                        for run in paragraph.runs:
                            set_arabic_font(run, font_size=14)

                # Bold header row
                if header_row and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

    return {
        "success": True,
        "doc_id": doc_id,
        "message": f"Table ({rows}x{cols}) added successfully"
    }


@mcp.tool()
def save_document(doc_id: str) -> dict:
    """
    Save the document to disk and return file information.

    Args:
        doc_id: Document ID to save

    Returns:
        dict with file_path, file_name, and download information
    """
    if doc_id not in active_documents:
        return {"success": False, "error": f"Document {doc_id} not found"}

    doc = active_documents[doc_id]
    metadata = document_metadata.get(doc_id, {})

    # Generate filename using tender_name from metadata (Arabic support)
    tender_name = metadata.get("tender_name", metadata.get("project_name", "مشروع"))
    safe_tender_name = "".join(c if c.isalnum() or c in (' ', '-', '_', 'ا', 'ب', 'ت', 'ث', 'ج', 'ح', 'خ', 'د', 'ذ', 'ر', 'ز', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ك', 'ل', 'م', 'ن', 'ه', 'و', 'ي') else '_' for c in tender_name)
    safe_tender_name = safe_tender_name.replace(' ', '_')[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Use standard naming format: RFP_{tender_name}_{doc_id}_{timestamp}.docx
    file_name = f"RFP_{safe_tender_name}_{doc_id[:8]}_{timestamp}.docx"
    file_path = DOCUMENTS_DIR / file_name

    # Save document
    doc.save(str(file_path))

    # Update metadata
    metadata["file_path"] = str(file_path)
    metadata["file_name"] = file_name
    metadata["saved_at"] = datetime.now().isoformat()

    # Construct download URL
    api_host = os.getenv("API_HOST", "http://localhost:7091")
    download_url = f"{api_host}/api/documents/download/{doc_id}"

    return {
        "success": True,
        "doc_id": doc_id,
        "file_name": file_name,
        "file_path": str(file_path),
        "download_url": download_url,
        "message": f"Document saved as {file_name}. Download at: {download_url}"
    }


@mcp.tool()
def get_document_preview(doc_id: str) -> dict:
    """
    Get a preview of the document structure and content.

    Args:
        doc_id: Document ID

    Returns:
        dict with document preview information
    """
    if doc_id not in active_documents:
        return {"success": False, "error": f"Document {doc_id} not found"}

    doc = active_documents[doc_id]
    metadata = document_metadata.get(doc_id, {})

    # Extract structure
    preview = {
        "success": True,
        "doc_id": doc_id,
        "title": metadata.get("title", "Untitled"),
        "project_name": metadata.get("project_name", ""),
        "created_at": metadata.get("created_at", ""),
        "sections": metadata.get("sections", []),
        "section_count": len(metadata.get("sections", [])),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables)
    }

    # Generate preview text
    preview_lines = [f"# {metadata.get('title', 'RFP Document')}"]
    preview_lines.append(f"\nProject: {metadata.get('project_name', 'N/A')}")
    preview_lines.append(f"Sections: {len(metadata.get('sections', []))}")
    preview_lines.append("\n## Document Structure:\n")

    for i, section in enumerate(metadata.get("sections", []), 1):
        indent = "  " * (section.get("level", 1) - 1)
        # Support both 'heading' (from add_section) and 'title' (from create_arabic_rfp_document)
        # Normalize to 'heading' for frontend consistency
        section_title = section.get('heading', section.get('title', 'Untitled'))
        section['heading'] = section_title  # Ensure 'heading' field exists for frontend
        preview_lines.append(f"{indent}{i}. {section_title}")

    preview["preview_text"] = "\n".join(preview_lines)

    return preview


@mcp.tool()
def list_documents() -> dict:
    """
    List all active documents in memory.

    Returns:
        dict with list of active documents
    """
    documents = []
    for doc_id, metadata in document_metadata.items():
        documents.append({
            "doc_id": doc_id,
            "title": metadata.get("title", "Untitled"),
            "project_name": metadata.get("project_name", ""),
            "created_at": metadata.get("created_at", ""),
            "sections": len(metadata.get("sections", []))
        })

    return {
        "success": True,
        "count": len(documents),
        "documents": documents
    }


@mcp.tool()
def delete_document(doc_id: str) -> dict:
    """
    Delete a document from memory and disk.

    Args:
        doc_id: Document ID to delete

    Returns:
        dict with success status
    """
    if doc_id not in active_documents:
        return {"success": False, "error": f"Document {doc_id} not found"}

    # Remove from memory
    del active_documents[doc_id]

    metadata = document_metadata.pop(doc_id, {})

    # Delete file if exists
    file_path = metadata.get("file_path")
    if file_path and Path(file_path).exists():
        Path(file_path).unlink()

    return {
        "success": True,
        "doc_id": doc_id,
        "message": "Document deleted successfully"
    }


@mcp.tool()
def list_reference_rfps() -> dict:
    """
    List all reference RFP documents available in the templates directory.

    These reference documents are examples of actual Saudi government RFPs
    that can be analyzed to understand structure, style, and content patterns.

    Returns:
        dict with list of available reference RFP files
    """
    reference_dir = Path("/app/application/templates/RFPs")

    if not reference_dir.exists():
        return {
            "success": False,
            "error": "Reference RFPs directory not found"
        }

    reference_files = []
    for file_path in reference_dir.glob("*"):
        if file_path.is_file() and file_path.suffix.lower() in ['.pdf', '.docx', '.doc']:
            reference_files.append({
                "name": file_path.name,
                "path": str(file_path),
                "type": file_path.suffix.lower(),
                "size_mb": round(file_path.stat().st_size / (1024 * 1024), 2)
            })

    return {
        "success": True,
        "count": len(reference_files),
        "reference_rfps": reference_files,
        "message": f"Found {len(reference_files)} reference RFP documents"
    }


@mcp.tool()
def extract_reference_rfp_content(file_name: str, max_paragraphs: int = 100) -> dict:
    """
    Extract text content from a reference RFP document (DOCX only for now).

    This tool allows the agent to read and analyze reference RFP documents
    to understand their structure, style, sections, and content patterns.

    Args:
        file_name: Name of the reference RFP file to extract
        max_paragraphs: Maximum number of paragraphs to extract (default 100)

    Returns:
        dict with extracted content, structure analysis, and sections
    """
    reference_dir = Path("/app/application/templates/RFPs")
    file_path = reference_dir / file_name

    if not file_path.exists():
        return {
            "success": False,
            "error": f"Reference RFP file '{file_name}' not found"
        }

    if file_path.suffix.lower() not in ['.docx']:
        return {
            "success": False,
            "error": "Currently only DOCX files are supported for content extraction. PDF support coming soon."
        }

    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))

        # Extract paragraphs
        paragraphs = []
        sections = []
        current_section = None

        for i, para in enumerate(doc.paragraphs):
            if i >= max_paragraphs:
                break

            text = para.text.strip()
            if not text:
                continue

            # Detect section headings (heuristic: short paragraphs with Arabic keywords)
            is_heading = False
            if len(text) < 150 and any(keyword in text for keyword in ["القسم", "الباب", "الفصل", "المادة"]):
                is_heading = True
                current_section = {
                    "heading": text,
                    "paragraph_index": i,
                    "content": []
                }
                sections.append(current_section)

            paragraphs.append({
                "index": i,
                "text": text,
                "is_heading": is_heading,
                "length": len(text)
            })

            if current_section and not is_heading:
                current_section["content"].append(text)

        # Extract tables
        tables_info = []
        for i, table in enumerate(doc.tables):
            if i >= 10:  # Limit to first 10 tables
                break

            rows = len(table.rows)
            cols = len(table.columns) if rows > 0 else 0

            # Extract header row if exists
            header = []
            if rows > 0:
                header = [cell.text.strip() for cell in table.rows[0].cells]

            tables_info.append({
                "table_index": i,
                "rows": rows,
                "columns": cols,
                "header": header[:5]  # First 5 columns only
            })

        # Analyze document structure
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)

        return {
            "success": True,
            "file_name": file_name,
            "structure": {
                "total_paragraphs": total_paragraphs,
                "extracted_paragraphs": len(paragraphs),
                "total_tables": total_tables,
                "detected_sections": len(sections)
            },
            "sections": sections[:20],  # First 20 sections
            "paragraphs": paragraphs[:50],  # First 50 paragraphs for context
            "tables": tables_info,
            "message": f"Extracted content from '{file_name}' - {len(sections)} sections, {len(paragraphs)} paragraphs, {len(tables_info)} tables"
        }

    except Exception as e:
        logger.error(f"Error extracting reference RFP content: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Failed to extract content: {str(e)}"
        }


@mcp.tool()
def analyze_reference_rfp_style(file_name: str) -> dict:
    """
    Analyze the writing style, tone, and patterns in a reference RFP document.

    This tool helps the agent understand how to write RFPs in a similar style
    by analyzing terminology, sentence structure, and content patterns.

    Args:
        file_name: Name of the reference RFP file to analyze

    Returns:
        dict with style analysis, common terms, and writing patterns
    """
    # First extract content
    content_result = extract_reference_rfp_content(file_name, max_paragraphs=50)

    if not content_result.get("success"):
        return content_result

    paragraphs = content_result.get("paragraphs", [])
    sections = content_result.get("sections", [])

    # Analyze common terms (simple frequency analysis)
    from collections import Counter

    all_text = " ".join([p["text"] for p in paragraphs])
    words = all_text.split()

    # Find common Arabic terms (simple heuristic)
    arabic_terms = [w for w in words if any('\u0600' <= c <= '\u06FF' for c in w)]
    common_terms = Counter(arabic_terms).most_common(30)

    # Analyze sentence patterns
    avg_paragraph_length = sum(p["length"] for p in paragraphs) / len(paragraphs) if paragraphs else 0

    # Detect tone indicators
    formal_indicators = ["يجب", "ينبغي", "يلتزم", "وفقاً", "حسب", "بموجب"]
    formal_count = sum(all_text.count(term) for term in formal_indicators)

    return {
        "success": True,
        "file_name": file_name,
        "style_analysis": {
            "avg_paragraph_length": round(avg_paragraph_length, 2),
            "total_words": len(words),
            "arabic_words": len(arabic_terms),
            "formality_score": formal_count,
            "tone": "formal" if formal_count > 10 else "moderate"
        },
        "common_terms": [{"term": term, "frequency": freq} for term, freq in common_terms[:15]],
        "section_patterns": [
            {"heading": s["heading"], "content_paragraphs": len(s["content"])}
            for s in sections[:10]
        ],
        "writing_patterns": {
            "uses_numbered_lists": "1." in all_text or "٠" in all_text,
            "uses_bullet_points": "•" in all_text or "-" in all_text,
            "includes_tables": len(content_result.get("tables", [])) > 0
        },
        "message": f"Analyzed writing style of '{file_name}'"
    }


if __name__ == "__main__":
    # Run the FastMCP server with HTTP transport
    import uvicorn

    host = os.getenv("MCP_SERVER_HOST", "0.0.0.0")
    port = int(os.getenv("MCP_SERVER_PORT", "8080"))

    print(f"Starting MCP-Doc Server on {host}:{port}")
    print(f"Documents will be saved to: {DOCUMENTS_DIR}")

    # Start FastMCP server
    mcp.run(transport="http", host=host, port=port)
