"""
DOCX Template Placeholder Extraction and Management Service
Handles parsing of DOCX templates to extract placeholders and dropdown fields
"""

import re
import logging
from typing import Dict, List, Set, Optional, Any, Tuple
from pathlib import Path
from dataclasses import dataclass, field
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass
class PlaceholderInfo:
    """Information about a placeholder in the document"""
    name: str
    placeholder_text: str  # Full placeholder with brackets {{name}}
    locations: List[str] = field(default_factory=list)  # Where it appears (paragraph, table, etc.)
    count: int = 0
    is_required: bool = True
    data_type: str = "text"  # text, number, date, dropdown, table
    validation_rules: Dict[str, Any] = field(default_factory=dict)
    description: str = ""
    special_instructions: str = ""


@dataclass
class DropdownField:
    """Information about a dropdown field in the document"""
    location: str  # paragraph or table cell reference
    text: str  # Original text (e.g., "اختيار عنصر")
    options: List[str] = field(default_factory=list)
    selected_value: Optional[str] = None


class DocxPlaceholderService:
    """Service for extracting and managing placeholders in DOCX templates"""

    # Regex pattern to match placeholders like {{placeholder_name}}
    PLACEHOLDER_PATTERN = r'\{\{([^}]+)\}\}'

    # Common dropdown indicators in Arabic RFP templates
    DROPDOWN_INDICATORS = [
        "اختيار عنصر",
        "اضغط هنا للاختيار",
        "اختر من القائمة",
        "حدد الخيار",
        "[اختيار]"
    ]

    # Special placeholders with specific generation rules
    SPECIAL_PLACEHOLDERS = {
        "project_scope": {
            "description": "نطاق العمل",
            "instructions": """يجب كتابة نطاق العمل بوضوح ودقة باللغة العربية.
            عدم الإشارة إلى منتج أو علامة تجارية محددة.
            تحديد مخرجات المشروع.
            متطلبات التدريب ونقل المعرفة إذا لزم الأمر."""
        },
        "work_program_phases": {
            "description": "مراحل البرنامج الزمني",
            "instructions": """تحديد مراحل التنفيذ ومدة كل مرحلة.
            المرحلة الأولى: [الوصف] مدة [المدة]
            المرحلة الثانية: [الوصف] مدة [المدة]"""
        },
        "work_program_payment_method": {
            "description": "طريقة الدفع",
            "instructions": """تحديد طريقة الدفع ونسب الدفعات.
            الدفعة الأولى: X% بعد [الحدث]
            الدفعة الثانية: X% بعد [الحدث]"""
        },
        "work_execution_method": {
            "description": "طريقة تنفيذ الأعمال",
            "instructions": """الخدمة التي سيتم عملها من قبل المتعاقد.
            التفاصيل المتعلقة بالخدمة.
            المواد التي سيتم استعمالها.
            القياسات المتعلقة بالمواد.
            تفاصيل الاختبارات المطلوبة."""
        }
    }

    def __init__(self, template_path: str):
        """Initialize the service with a template file path"""
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")

        self.document: Optional[Document] = None
        self.placeholders: Dict[str, PlaceholderInfo] = {}
        self.dropdown_fields: List[DropdownField] = []

    def load_template(self) -> Document:
        """Load the DOCX template file"""
        try:
            self.document = Document(self.template_path)
            logger.info(f"Successfully loaded template: {self.template_path}")
            return self.document
        except Exception as e:
            logger.error(f"Failed to load template: {e}")
            raise

    def extract_placeholders(self) -> Dict[str, PlaceholderInfo]:
        """Extract all placeholders from the document"""
        if not self.document:
            self.load_template()

        self.placeholders = {}

        # Extract from paragraphs
        for i, paragraph in enumerate(self.document.paragraphs):
            self._extract_from_text(paragraph.text, f"paragraph_{i}")

        # Extract from tables
        for t_idx, table in enumerate(self.document.tables):
            self._extract_from_table(table, f"table_{t_idx}")

        # Extract from headers and footers
        for section in self.document.sections:
            # Header
            if section.header:
                for p_idx, paragraph in enumerate(section.header.paragraphs):
                    self._extract_from_text(paragraph.text, f"header_{p_idx}")

            # Footer
            if section.footer:
                for p_idx, paragraph in enumerate(section.footer.paragraphs):
                    self._extract_from_text(paragraph.text, f"footer_{p_idx}")

        # Add special instructions for known placeholders
        for placeholder_name, info in self.placeholders.items():
            if placeholder_name in self.SPECIAL_PLACEHOLDERS:
                special_info = self.SPECIAL_PLACEHOLDERS[placeholder_name]
                info.description = special_info["description"]
                info.special_instructions = special_info["instructions"]

        logger.info(f"Extracted {len(self.placeholders)} unique placeholders")
        return self.placeholders

    def _extract_from_text(self, text: str, location: str):
        """Extract placeholders from a text string"""
        if not text:
            return

        matches = re.finditer(self.PLACEHOLDER_PATTERN, text)
        for match in matches:
            placeholder_name = match.group(1).strip()
            placeholder_full = match.group(0)

            if placeholder_name not in self.placeholders:
                self.placeholders[placeholder_name] = PlaceholderInfo(
                    name=placeholder_name,
                    placeholder_text=placeholder_full,
                    locations=[location],
                    count=1
                )
            else:
                self.placeholders[placeholder_name].locations.append(location)
                self.placeholders[placeholder_name].count += 1

    def _extract_from_table(self, table: Table, table_id: str):
        """Extract placeholders from a table"""
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_location = f"{table_id}_r{r_idx}_c{c_idx}"

                # Extract from cell paragraphs
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    self._extract_from_text(
                        paragraph.text,
                        f"{cell_location}_p{p_idx}"
                    )

    def extract_dropdown_fields(self) -> List[DropdownField]:
        """Extract dropdown fields from the document"""
        if not self.document:
            self.load_template()

        self.dropdown_fields = []

        # Check paragraphs
        for i, paragraph in enumerate(self.document.paragraphs):
            self._check_for_dropdown(paragraph.text, f"paragraph_{i}")

        # Check tables
        for t_idx, table in enumerate(self.document.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_location = f"table_{t_idx}_r{r_idx}_c{c_idx}"
                    for paragraph in cell.paragraphs:
                        self._check_for_dropdown(paragraph.text, cell_location)

        logger.info(f"Found {len(self.dropdown_fields)} dropdown fields")
        return self.dropdown_fields

    def _check_for_dropdown(self, text: str, location: str):
        """Check if text contains dropdown indicators"""
        if not text:
            return

        for indicator in self.DROPDOWN_INDICATORS:
            if indicator in text:
                # Try to extract options if they're listed nearby
                options = self._extract_dropdown_options(text)

                self.dropdown_fields.append(DropdownField(
                    location=location,
                    text=indicator,
                    options=options
                ))

    def _extract_dropdown_options(self, text: str) -> List[str]:
        """Try to extract dropdown options from surrounding text"""
        options = []

        # Look for common patterns like:
        # (خيار1، خيار2، خيار3)
        # [خيار1 | خيار2 | خيار3]

        # Pattern 1: Options in parentheses separated by commas
        pattern1 = r'\(([\u0600-\u06FF\s,،/]+)\)'
        match = re.search(pattern1, text)
        if match:
            options_text = match.group(1)
            options = [opt.strip() for opt in re.split(r'[,،/]', options_text) if opt.strip()]

        # Pattern 2: Options in brackets separated by pipes
        if not options:
            pattern2 = r'\[([\u0600-\u06FF\s|/]+)\]'
            match = re.search(pattern2, text)
            if match:
                options_text = match.group(1)
                options = [opt.strip() for opt in re.split(r'[|/]', options_text) if opt.strip()]

        return options

    def get_required_placeholders(self) -> List[str]:
        """Get list of required placeholder names"""
        return [
            name for name, info in self.placeholders.items()
            if info.is_required
        ]

    def get_placeholder_summary(self) -> Dict[str, Any]:
        """Get a summary of all placeholders and their information"""
        summary = {
            "total_placeholders": len(self.placeholders),
            "required_placeholders": len(self.get_required_placeholders()),
            "dropdown_fields": len(self.dropdown_fields),
            "placeholders": {}
        }

        for name, info in self.placeholders.items():
            summary["placeholders"][name] = {
                "count": info.count,
                "locations": info.locations,
                "is_required": info.is_required,
                "description": info.description,
                "has_special_instructions": bool(info.special_instructions)
            }

        return summary

    def validate_placeholder_data(self, data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """
        Validate that all required placeholders have data
        Returns: (is_valid, list_of_missing_placeholders)
        """
        missing = []

        for placeholder_name in self.get_required_placeholders():
            if placeholder_name not in data or not data[placeholder_name]:
                missing.append(placeholder_name)

        return len(missing) == 0, missing

    def get_placeholder_instructions(self, placeholder_name: str) -> Optional[str]:
        """Get special instructions for a specific placeholder"""
        if placeholder_name in self.placeholders:
            return self.placeholders[placeholder_name].special_instructions
        return None


# Utility function for quick extraction
def extract_template_placeholders(template_path: str) -> Dict[str, Any]:
    """Quick utility to extract all placeholders from a template"""
    service = DocxPlaceholderService(template_path)
    service.extract_placeholders()
    service.extract_dropdown_fields()
    return service.get_placeholder_summary()


if __name__ == "__main__":
    # Test the service with the actual template
    template_path = "inputs/templates/rfp_template_with_placeholders.docx"

    try:
        service = DocxPlaceholderService(template_path)
        placeholders = service.extract_placeholders()
        dropdowns = service.extract_dropdown_fields()

        print(f"Found {len(placeholders)} unique placeholders:")
        for name, info in placeholders.items():
            print(f"  - {{{{name}}}}: appears {info.count} times")
            if info.description:
                print(f"    Description: {info.description}")

        print(f"\nFound {len(dropdowns)} dropdown fields:")
        for dropdown in dropdowns:
            print(f"  - Location: {dropdown.location}")
            print(f"    Text: {dropdown.text}")
            if dropdown.options:
                print(f"    Options: {', '.join(dropdown.options)}")

        print("\nSummary:")
        print(service.get_placeholder_summary())

    except Exception as e:
        print(f"Error: {e}")