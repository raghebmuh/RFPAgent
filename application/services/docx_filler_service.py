"""
DOCX Template Filler Service
Fills placeholders in DOCX templates while preserving formatting and structure
"""

import os
import logging
import re
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
import arabic_reshaper
from bidi.algorithm import get_display

from application.services.rfp_content_generator import RFPContentGenerator

logger = logging.getLogger(__name__)


class DocxFillerService:
    """
    Service for filling DOCX template placeholders while preserving formatting
    """

    PLACEHOLDER_PATTERN = r'\{\{([^}]+)\}\}'
    DROPDOWN_INDICATORS = [
        "اختيار عنصر",
        "اضغط هنا للاختيار",
        "اختر من القائمة",
        "[اختيار]"
    ]

    def __init__(self, template_path: str):
        """Initialize with template file path"""
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")

        self.document: Optional[Document] = None
        self.content_generator = RFPContentGenerator()

    def fill_template(self, placeholder_data: Dict[str, Any], output_path: str) -> str:
        """
        Fill the template with provided data and save to output path
        Returns the path of the generated file
        """
        try:
            # Load the template
            self.document = Document(self.template_path)
            logger.info(f"Loaded template from {self.template_path}")

            # Generate content for special placeholders
            enriched_data = self._enrich_placeholder_data(placeholder_data)

            # Process all paragraphs
            for paragraph in self.document.paragraphs:
                self._process_paragraph(paragraph, enriched_data)

            # Process all tables
            for table in self.document.tables:
                self._process_table(table, enriched_data)

            # Process headers and footers
            for section in self.document.sections:
                # Process header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self._process_paragraph(paragraph, enriched_data)

                # Process footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self._process_paragraph(paragraph, enriched_data)

            # Save the filled document
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(output_file)

            logger.info(f"Document saved to {output_file}")
            return str(output_file)

        except Exception as e:
            logger.error(f"Error filling template: {e}")
            raise

    def _enrich_placeholder_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Enrich placeholder data by generating content for special placeholders
        """
        enriched = data.copy()

        # Generate content for special placeholders using content generator
        special_placeholders = [
            "project_scope",
            "work_program_phases",
            "work_program_payment_method",
            "work_execution_method",
            "evaluation_criteria",
            "required_certificates",
            "technical_specifications",
            "quality_standards",
            "safety_requirements",
            "deliverables",
            "project_objectives"
        ]

        for placeholder in special_placeholders:
            # Only generate if not already provided or if provided value is too short
            if placeholder not in enriched or len(str(enriched.get(placeholder, ""))) < 50:
                enriched[placeholder] = self.content_generator.generate_content(placeholder, data)

        # Add metadata
        enriched["generation_date"] = datetime.now().strftime("%Y/%m/%d")
        enriched["generation_time"] = datetime.now().strftime("%H:%M")

        # Add default values for missing optional fields
        defaults = {
            "classification": "غير محدد",
            "warranty_period": "12 شهر",
            "local_content_percentage": "30",
            "technical_weight": "60",
            "financial_weight": "40"
        }

        for key, default_value in defaults.items():
            if key not in enriched or not enriched[key]:
                enriched[key] = default_value

        return enriched

    def _process_paragraph(self, paragraph: Paragraph, data: Dict[str, Any]):
        """
        Process a paragraph to replace placeholders and dropdowns
        """
        if not paragraph.text:
            return

        original_text = paragraph.text

        # Check for dropdown indicators first
        for dropdown_indicator in self.DROPDOWN_INDICATORS:
            if dropdown_indicator in original_text:
                # Replace dropdown with appropriate selection
                replacement = self._get_dropdown_replacement(original_text, data)
                self._replace_paragraph_text(paragraph, dropdown_indicator, replacement)

        # Replace placeholders
        placeholders = re.finditer(self.PLACEHOLDER_PATTERN, paragraph.text)
        replacements = []

        for match in placeholders:
            placeholder_name = match.group(1).strip()
            placeholder_full = match.group(0)

            if placeholder_name in data:
                replacement_value = str(data[placeholder_name])
                replacements.append((placeholder_full, replacement_value))

        # Apply replacements
        for old, new in replacements:
            self._replace_paragraph_text(paragraph, old, new)

    def _process_table(self, table: Table, data: Dict[str, Any]):
        """
        Process all cells in a table
        """
        for row in table.rows:
            for cell in row.cells:
                self._process_cell(cell, data)

    def _process_cell(self, cell: _Cell, data: Dict[str, Any]):
        """
        Process a table cell to replace placeholders
        """
        for paragraph in cell.paragraphs:
            self._process_paragraph(paragraph, data)

    def _replace_paragraph_text(self, paragraph: Paragraph, old_text: str, new_text: str):
        """
        Replace text in a paragraph while preserving formatting
        """
        # Check if we need Arabic text reshaping
        if self._is_arabic(new_text):
            # Apply Arabic reshaping for proper display
            # Note: python-docx usually handles RTL correctly, but we reshape for safety
            new_text = self._reshape_arabic_text(new_text)

        # Simple approach: Replace in the paragraph text directly
        # This preserves basic formatting but might lose complex run-level formatting
        if old_text in paragraph.text:
            # Store the alignment
            original_alignment = paragraph.alignment

            # Get all runs and their properties
            runs_data = []
            for run in paragraph.runs:
                runs_data.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size
                })

            # Combine text from all runs
            full_text = ''.join(run['text'] for run in runs_data)

            # Replace the placeholder
            new_full_text = full_text.replace(old_text, new_text)

            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""

            # Add new text with formatting
            # Try to maintain formatting from the first run
            if runs_data:
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = new_full_text

                # Apply formatting from original first run
                if runs_data[0]['bold'] is not None:
                    run.bold = runs_data[0]['bold']
                if runs_data[0]['italic'] is not None:
                    run.italic = runs_data[0]['italic']
                if runs_data[0]['underline'] is not None:
                    run.underline = runs_data[0]['underline']
                if runs_data[0]['font_name']:
                    run.font.name = runs_data[0]['font_name']
                if runs_data[0]['font_size']:
                    run.font.size = runs_data[0]['font_size']
            else:
                paragraph.text = new_full_text

            # Restore alignment
            if original_alignment:
                paragraph.alignment = original_alignment

            # Set RTL for Arabic text
            if self._is_arabic(new_text):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _get_dropdown_replacement(self, text: str, data: Dict[str, Any]) -> str:
        """
        Determine appropriate replacement for dropdown fields
        """
        # Context-based replacement
        text_lower = text.lower()

        # Project type dropdown
        if "نوع" in text or "المشروع" in text:
            return data.get("project_type", "تقنية المعلومات")

        # Payment method dropdown
        elif "الدفع" in text or "طريقة" in text:
            return data.get("payment_method", "دفعات حسب المراحل")

        # Duration dropdown
        elif "المدة" in text or "فترة" in text:
            duration = data.get("duration_months", 6)
            return f"{duration} شهر"

        # Training required dropdown
        elif "التدريب" in text:
            return data.get("training_required", "نعم")

        # Default replacement
        return "حسب المتطلبات"

    def _is_arabic(self, text: str) -> bool:
        """
        Check if text contains Arabic characters
        """
        arabic_pattern = r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]'
        return bool(re.search(arabic_pattern, text))

    def _reshape_arabic_text(self, text: str) -> str:
        """
        Reshape Arabic text for proper display
        """
        try:
            # Apply Arabic reshaping
            reshaped = arabic_reshaper.reshape(text)
            # Apply bidi algorithm for mixed text
            # Note: We don't apply get_display here as DOCX handles RTL
            return reshaped
        except Exception as e:
            logger.warning(f"Error reshaping Arabic text: {e}")
            return text

    def extract_filled_content(self, data: Dict[str, Any]) -> Dict[str, str]:
        """
        Extract the filled content for preview purposes
        Returns a dictionary of placeholder names to filled content
        """
        enriched_data = self._enrich_placeholder_data(data)
        filled_content = {}

        for key, value in enriched_data.items():
            if isinstance(value, str) and len(value) > 0:
                filled_content[key] = value

        return filled_content

    def generate_preview_text(self, data: Dict[str, Any]) -> str:
        """
        Generate a text preview of the filled document
        """
        filled_content = self.extract_filled_content(data)
        preview = "معاينة وثيقة طلب تقديم العروض\n" + "="*50 + "\n\n"

        # Key sections to preview
        key_sections = [
            ("معلومات المشروع", ["entity_name", "project_name", "tender_number", "project_type"]),
            ("نطاق العمل", ["project_scope"]),
            ("البرنامج الزمني", ["duration_months", "work_program_phases"]),
            ("طريقة الدفع", ["work_program_payment_method"]),
            ("طريقة التنفيذ", ["work_execution_method"]),
            ("معايير التقييم", ["evaluation_criteria"])
        ]

        for section_title, fields in key_sections:
            preview += f"\n{section_title}:\n" + "-"*30 + "\n"
            for field in fields:
                if field in filled_content and filled_content[field]:
                    value = filled_content[field]
                    # Truncate long values for preview
                    if len(value) > 500:
                        value = value[:500] + "..."
                    preview += f"{value}\n"

        preview += "\n" + "="*50 + "\n"
        preview += "الوثيقة جاهزة للتحميل بصيغة DOCX قابلة للتعديل"

        return preview

    def get_document_sections(self) -> List[Dict[str, Any]]:
        """
        Extract document sections structure for display
        """
        if not self.document:
            self.document = Document(self.template_path)

        sections = []
        current_section = None
        section_level = 1

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()

            # Detect section headers (القسم الأول، القسم الثاني، etc.)
            if text and ("القسم" in text or "الفصل" in text):
                if current_section:
                    sections.append(current_section)

                current_section = {
                    "heading": text,
                    "level": section_level,
                    "subsections": []
                }

            # Detect subsections (numbers or bullets)
            elif text and current_section and (
                text[0].isdigit() or
                text.startswith("•") or
                text.startswith("-") or
                text.startswith("*")
            ):
                current_section["subsections"].append({
                    "heading": text,
                    "level": section_level + 1
                })

        # Add the last section
        if current_section:
            sections.append(current_section)

        return sections


# Utility function for quick document generation
def fill_rfp_template(
    template_path: str,
    placeholder_data: Dict[str, Any],
    output_path: str
) -> str:
    """
    Quick utility to fill an RFP template
    Returns the path of the generated document
    """
    service = DocxFillerService(template_path)
    return service.fill_template(placeholder_data, output_path)


# Preview generation utility
def generate_rfp_preview(template_path: str, placeholder_data: Dict[str, Any]) -> str:
    """Generate a text preview of the filled RFP document"""
    service = DocxFillerService(template_path)
    return service.generate_preview_text(placeholder_data)